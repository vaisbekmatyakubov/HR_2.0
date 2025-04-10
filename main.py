# -*- coding: utf-8 -*-
import streamlit as st
from docx import Document
from io import BytesIO
import datetime
from babel.dates import format_date
import locale

# --- Streamlit Саҳифа Конфигурацияси ---
st.set_page_config(layout="wide", page_title="Хизмат Сафари Ҳисоботи Генератори")

# --- Локални ўрнатиш ---
try:
    locale.setlocale(locale.LC_TIME, 'uz_UZ.UTF-8') # Linux/macOS
    BABEL_LOCALE = 'uz_UZ_cyrl'
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, 'uz_UZ_cyrl.UTF-8') # Баъзи тизимлар учун
        BABEL_LOCALE = 'uz_UZ_cyrl'
    except locale.Error:
        BABEL_LOCALE = 'en_US' # Захира локал

# --- Шаблонни тўлдириш функцияси ---
def generate_report(template_path, output_name, data):
    """
    Word шаблонини берилган маълумотлар билан тўлдиради ва
    хотирадаги файл буферини қайтаради.
    """
    try:
        doc = Document(template_path)
    except Exception as e:
        st.error(f"Шаблон файлини ўқишда хатолик ({template_path}): {e}")
        st.error("Ишонч ҳосил қилингки, файл мавжуд ва тўғри форматда.")
        return None

    def format_uzbek_date(date_obj):
        if date_obj:
            try:
                return format_date(date_obj, format='d MMMM<y_bin_46> \'йил\'', locale=BABEL_LOCALE)
            except Exception as e:
                return date_obj.strftime('%d.%m.%Y')
        return ""

    buyruq_sanasi_str = format_uzbek_date(data['buyruq_sanasi'])
    safar_boshlanish_str = format_uzbek_date(data['safar_boshlanish'])
    safar_tugash_str = format_uzbek_date(data['safar_tugash'])
    hisobot_sanasi_str = format_uzbek_date(data['hisobot_sanasi'])

    buyruq_r_s = f"{data['buyruq_raqami']}-сонли {buyruq_sanasi_str}даги" if data['buyruq_raqami'] and data['buyruq_sanasi'] else ""
    sanalar = f"{safar_boshlanish_str} дан {safar_tugash_str} гача" if data['safar_boshlanish'] and data['safar_tugash'] else ""
    maqsadlar_text = "\n".join([f"– {m}" for m in data['maqsadlar'] if m.strip()])
    natijalar_text = "\n".join([f"– {n}" for n in data['safar_natijalari'] if n.strip()])

    kunlik_hisobot_text = ""
    for i, report in enumerate(data['kunlik_hisobotlar']):
        kun_nomi = report.get('kun_nomi', f"{i+1}-кун")
        kun_sanasi_str = format_uzbek_date(report.get('kun_sanasi'))
        tavsif = report.get('tavsif', '')
        if tavsif.strip():
             kunlik_hisobot_text += f"{kun_nomi} ({kun_sanasi_str}):\n{tavsif}\n\n"

    replacements = {
        "{{LAVOZIM_FIO}}": f"{data['lavozim']} {data['fio']}",
        "{{MANZIL}}": data['manzil'],
        "{{BUYRUQ_R_S}}": buyruq_r_s,
        "{{SANALAR}}": sanalar,
        "{{MAQSADLAR}}": maqsadlar_text,
        "{{UMUMIY_QISM}}": data['umumiy_qism'], # ЯНГИ: Умумий қисм қўшилди
        "{{KUNLIK_HISOBOT}}": kunlik_hisobot_text.strip(),
        "{{SAFAR_NATIJALARI}}": natijalar_text,
        "{{BOLIM_BOSHLIGI_LAVOZIMI}}": data['bolim_boshligi_lavozimi'],
        "{{BOLIM_BOSHLIGI_FIO}}": data['bolim_boshligi_fio'],
        "{{HISOBOT_SANASI}}": hisobot_sanasi_str,
    }

    # Placeholder'ларни алмаштириш
    for p in doc.paragraphs:
        # Бутун параграф placeholder'га тенг бўлса (масалан, {{MAQSADLAR}}, {{UMUMIY_QISM}})
        if p.text.strip() in replacements:
             key = p.text.strip()
             p.text = str(replacements[key])
             # Эски run'ларни тозалаш (формат сақланмаслиги мумкин)
             for run in p.runs[1:]: run.clear()
             if p.runs: p.runs[0].text = str(replacements[key])
        else:
             # Placeholder параграф ичида бўлса
             inline = p.runs
             full_text = "".join(run.text for run in inline)
             changed = False
             for key, value in replacements.items():
                 if key in full_text:
                     full_text = full_text.replace(key, str(value))
                     changed = True
             if changed:
                  for i in range(len(p.runs)): p.runs[i].text = ''
                  if p.runs: p.runs[0].text = full_text
                  else: p.add_run(full_text)

    # Жадвалларда алмаштириш
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip() in replacements:
                         key = p.text.strip()
                         p.text = str(replacements[key])
                         for run in p.runs[1:]: run.clear()
                         if p.runs: p.runs[0].text = str(replacements[key])
                    else:
                         inline = p.runs
                         full_text = "".join(run.text for run in inline)
                         changed = False
                         for key, value in replacements.items():
                             if key in full_text:
                                 full_text = full_text.replace(key, str(value))
                                 changed = True
                         if changed:
                             for i in range(len(p.runs)): p.runs[i].text = ''
                             if p.runs: p.runs[0].text = full_text
                             else: p.add_run(full_text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Streamlit Интерфейси ---
st.title("📄 Хизмат Сафари Ҳисоботи Авто-тўлдирувчиси")

# --- Сессия ҳолатини инициализация қилиш ---
if 'maqsadlar' not in st.session_state:
    st.session_state.maqsadlar = [""]
if 'safar_natijalari' not in st.session_state:
    st.session_state.safar_natijalari = [""]
if 'kunlik_hisobotlar' not in st.session_state:
    st.session_state.kunlik_hisobotlar = [{'kun_nomi': 'Биринчи кун', 'kun_sanasi': None, 'tavsif': ''}]
# Асосий майдонлар учун калитлар (формасиз ишлаш учун)
if 'main_fio' not in st.session_state: st.session_state.main_fio = ""
if 'main_lavozim' not in st.session_state: st.session_state.main_lavozim = ""
if 'main_manzil' not in st.session_state: st.session_state.main_manzil = ""
if 'main_buyruq_raqami' not in st.session_state: st.session_state.main_buyruq_raqami = ""
if 'main_buyruq_sanasi' not in st.session_state: st.session_state.main_buyruq_sanasi = None
if 'main_safar_boshlanish' not in st.session_state: st.session_state.main_safar_boshlanish = None
if 'main_safar_tugash' not in st.session_state: st.session_state.main_safar_tugash = None
if 'main_bolim_boshligi_lavozimi' not in st.session_state: st.session_state.main_bolim_boshligi_lavozimi = ""
if 'main_bolim_boshligi_fio' not in st.session_state: st.session_state.main_bolim_boshligi_fio = ""
if 'main_hisobot_sanasi' not in st.session_state: st.session_state.main_hisobot_sanasi = datetime.date.today()
if 'main_umumiy_qism' not in st.session_state: st.session_state.main_umumiy_qism = "" # ЯНГИ: Умумий қисм учун


# --- Файл номлари ---
TEMPLATE_FILE = 'template_modified.docx'
DEFAULT_OUTPUT_FILENAME = "Tayyor_Hisobot.docx"

# --- Асосий Маълумотлар (Формасиз) ---
st.subheader("📝 Сафар Маълумотлари ва Тасдиқловчи")
col1, col2 = st.columns(2)
with col1:
    st.text_input("Ф.И.О.", placeholder="Исмоилов Анвар Каримович", key="main_fio")
    st.text_input("Лавозими", placeholder="Бош мутахассис", key="main_lavozim")
    st.text_input("Сафар манзили (Шаҳар/Вилоят)", placeholder="Тошкент шаҳри", key="main_manzil")
    st.text_input("Ички буйруқ рақами", placeholder="55-Х/С", key="main_buyruq_raqami")
    st.date_input("Буйруқ санаси", value=st.session_state.main_buyruq_sanasi, format="DD.MM.YYYY", key="main_buyruq_sanasi")

with col2:
    st.date_input("Сафар бошланиш санаси", value=st.session_state.main_safar_boshlanish, format="DD.MM.YYYY", key="main_safar_boshlanish")
    st.date_input("Сафар тугаш санаси", value=st.session_state.main_safar_tugash, format="DD.MM.YYYY", key="main_safar_tugash")
    st.markdown("---")
    st.text_input("Бўлим бошлиғи лавозими", placeholder="Департамент директори", key="main_bolim_boshligi_lavozimi")
    st.text_input("Бўлим бошлиғи Ф.И.О.", placeholder="Ахмедов Бахтиёр Эркинович", key="main_bolim_boshligi_fio")
    st.date_input("Ҳисобот санаси", value=st.session_state.main_hisobot_sanasi, format="DD.MM.YYYY", key="main_hisobot_sanasi")

st.divider()

# --- ЯНГИ: Умумий қисм учун матн майдони ---
st.subheader("📄 Умумий қисм")
st.text_area(
    "Сафарнинг асосий сабаблари, унгача бўлган тафсилотлар ва омиллар ҳақида ёзинг.",
    key="main_umumiy_qism",
    height=150
)

st.divider()

# --- Динамик Мақсадлар ---
st.subheader("🎯 Хизмат сафаридан асосий мақсадлар")
maqsadlar_container = st.container()
maqsadlar_to_remove = []
for i in range(len(st.session_state.maqsadlar)):
    row = maqsadlar_container.columns([0.9, 0.1])
    st.session_state.maqsadlar[i] = row[0].text_area(
        f"Мақсад {i+1}", value=st.session_state.maqsadlar[i], key=f"maqsad_{i}",
        height=50, label_visibility="collapsed", placeholder=f"{i+1}-мақсадни киритинг..."
    )
    if row[1].button("❌", key=f"remove_maqsad_{i}", help="Ушбу мақсадни ўчириш"):
         maqsadlar_to_remove.append(i)

if maqsadlar_to_remove:
    for index in sorted(maqsadlar_to_remove, reverse=True):
         del st.session_state.maqsadlar[index]
    st.rerun()

if st.button("➕ Мақсад қўшиш", key="add_maqsad"):
    st.session_state.maqsadlar.append("")
    st.rerun()

st.divider()

# --- Динамик Кунлик Ҳисоботлар ---
st.subheader("🗓️ Хизмат сафари давомида амалга оширилган ишлар (кунлар бўйича)")
hisobot_container = st.container()
hisobotlar_to_remove = []

for i in range(len(st.session_state.kunlik_hisobotlar)):
    with hisobot_container.container(border=True):
         cols_hisobot_header = st.columns([0.85, 0.15])
         with cols_hisobot_header[0]:
             cols_hisobot = st.columns([1, 1])
             with cols_hisobot[0]:
                 st.session_state.kunlik_hisobotlar[i]['kun_nomi'] = st.text_input(
                     "Кун номи", value=st.session_state.kunlik_hisobotlar[i].get('kun_nomi', f'{i+1}-кун'),
                     key=f"kun_nomi_{i}", label_visibility="collapsed", placeholder=f"{i+1}-кун номи"
                 )
             with cols_hisobot[1]:
                 st.session_state.kunlik_hisobotlar[i]['kun_sanasi'] = st.date_input(
                     "Сана", value=st.session_state.kunlik_hisobotlar[i].get('kun_sanasi'),
                     key=f"kun_sana_{i}", format="DD.MM.YYYY", label_visibility="collapsed"
                 )

         if cols_hisobot_header[1].button("❌ Ўчириш", key=f"remove_hisobot_{i}", help="Ушбу кун ҳисоботини ўчириш"):
             hisobotlar_to_remove.append(i)

         st.session_state.kunlik_hisobotlar[i]['tavsif'] = st.text_area(
             "Бажарилган ишлар тавсифи", value=st.session_state.kunlik_hisobotlar[i].get('tavsif', ''),
             key=f"kun_tavsif_{i}", height=100, label_visibility="collapsed", placeholder="Бажарилган ишлар тавсифи..."
         )

if hisobotlar_to_remove:
    for index in sorted(hisobotlar_to_remove, reverse=True):
         del st.session_state.kunlik_hisobotlar[index]
    st.rerun()

if st.button("➕ Кун қўшиш", key="add_kun"):
     next_day_num = len(st.session_state.kunlik_hisobotlar) + 1
     kun_nomi_default = f"{next_day_num}-кун"
     last_date = None
     if st.session_state.kunlik_hisobotlar:
         valid_dates = [entry.get('kun_sanasi') for entry in st.session_state.kunlik_hisobotlar if entry.get('kun_sanasi') is not None]
         if valid_dates:
              last_entry_date = max(valid_dates)
              try: last_date = last_entry_date + datetime.timedelta(days=1)
              except TypeError: last_date = None
     st.session_state.kunlik_hisobotlar.append({'kun_nomi': kun_nomi_default, 'kun_sanasi': last_date, 'tavsif': ''})
     st.rerun()

st.divider()

# --- Динамик Сафар Натижалари ---
st.subheader("📊 Сафар натижалари")
natijalar_container = st.container()
natijalar_to_remove = []
for i in range(len(st.session_state.safar_natijalari)):
    row = natijalar_container.columns([0.9, 0.1])
    st.session_state.safar_natijalari[i] = row[0].text_area(
        f"Натижа {i+1}", value=st.session_state.safar_natijalari[i], key=f"natija_{i}",
        height=50, label_visibility="collapsed", placeholder=f"{i+1}-натижани киритинг..."
    )
    if row[1].button("❌", key=f"remove_natija_{i}", help="Ушбу натижани ўчириш"):
         natijalar_to_remove.append(i)

if natijalar_to_remove:
    for index in sorted(natijalar_to_remove, reverse=True):
         del st.session_state.safar_natijalari[index]
    st.rerun()

if st.button("➕ Натижа қўшиш", key="add_natija"):
    st.session_state.safar_natijalari.append("")
    st.rerun()

st.divider()

# --- Ҳисоботни Яратиш Тугмаси ---
generate_button = st.button("✅ Ҳисоботни Яратиш", key="generate_report_main", type="primary")

# --- Ҳисоботни яратиш ва юклаб олиш логикаси ---
if generate_button:
    # Маълумотларни йиғиш (st.session_state дан КАЛИТЛАР орқали)
    report_data = {
        'fio': st.session_state.get('main_fio', ''),
        'lavozim': st.session_state.get('main_lavozim', ''),
        'manzil': st.session_state.get('main_manzil', ''),
        'buyruq_raqami': st.session_state.get('main_buyruq_raqami', ''),
        'buyruq_sanasi': st.session_state.get('main_buyruq_sanasi'),
        'safar_boshlanish': st.session_state.get('main_safar_boshlanish'),
        'safar_tugash': st.session_state.get('main_safar_tugash'),
        'umumiy_qism': st.session_state.get('main_umumiy_qism', ''), # ЯНГИ: Умумий қисм қиймати
        'bolim_boshligi_lavozimi': st.session_state.get('main_bolim_boshligi_lavozimi', ''),
        'bolim_boshligi_fio': st.session_state.get('main_bolim_boshligi_fio', ''),
        'hisobot_sanasi': st.session_state.get('main_hisobot_sanasi'),
        # Динамик рўйхатлар
        'maqsadlar': [m for m in st.session_state.get('maqsadlar', []) if m.strip()],
        'safar_natijalari': [n for n in st.session_state.get('safar_natijalari', []) if n.strip()],
        'kunlik_hisobotlar': [k for k in st.session_state.get('kunlik_hisobotlar', []) if k.get('tavsif','').strip()],
    }

    fio_for_filename = st.session_state.get('main_fio', 'Noma\'lum')
    output_filename = f"Hisobot_{fio_for_filename.replace(' ', '_')}_{datetime.datetime.now().strftime('%Y%m%d')}.docx"

    st.info("⏳ Ҳисобот яратилмоқда, илтимос кутинг...")
    generated_file_buffer = generate_report(TEMPLATE_FILE, output_filename, report_data)

    if generated_file_buffer:
        st.success("🎉 Ҳисобот муваффақиятли яратилди!")
        st.download_button(
            label="📥 Ҳисоботни Юклаб Олиш (.docx)",
            data=generated_file_buffer,
            file_name=output_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("❌ Ҳисоботни яратишда хатолик юз берди. Консолдаги хабарларни ва шаблон файлини текширинг.")

